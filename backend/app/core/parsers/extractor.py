"""Resume document text extraction: PDF and DOCX.

PDF extraction uses a 5-strategy fallback chain so that modern resumes created
by Canva, Novoresume, Google Docs, Microsoft Word, Adobe Acrobat, and similar
tools are all handled — regardless of font encoding, XRef compression, or
embedded-font issues.

Strategy order for PDFs:
  1. pymupdf  get_text("text")   — fastest, works for most standard PDFs
  2. pymupdf  get_text("blocks") — catches some embedded-font / complex layout PDFs
  3. pymupdf  rawdict span walk  — catches PDFs where block extraction misses spans
  4. pymupdf  html + strip tags  — last-resort within pymupdf
  5. pdfminer.six                — completely independent PDF parser; catches PDFs
                                   with CID-keyed/Type3 fonts where PyMuPDF returns
                                   empty glyphs (common in Canva / Novoresume exports)
"""

from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class UnsupportedFileError(ValueError):
    pass


# ─────────────────────────────────────────────────────────────────────────────
# PDF helpers
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_pymupdf(data: bytes) -> str:
    """Try all four PyMuPDF extraction modes; return the best non-empty result."""
    try:
        import pymupdf as fitz  # new import name (PyMuPDF ≥ 1.24)
    except ImportError:
        import fitz  # legacy import — suppress deprecation via env or warnings filter
        import warnings
        warnings.filterwarnings("ignore", message=".*fitz API is deprecated.*")

    doc = fitz.open(stream=data, filetype="pdf")
    page_texts: list[str] = []

    for page in doc:
        text = ""

        # 1. plain text
        try:
            text = (page.get_text("text") or "").strip()
        except Exception as exc:
            logger.debug("pymupdf get_text('text') page %s failed: %s", page.number, exc)

        # 2. blocks
        if not text:
            try:
                blocks = page.get_text("blocks") or []
                text = "\n".join(
                    b[4] for b in blocks
                    if isinstance(b, (list, tuple)) and len(b) > 4 and b[4]
                ).strip()
            except Exception as exc:
                logger.debug("pymupdf get_text('blocks') page %s failed: %s", page.number, exc)

        # 3. rawdict span walk
        if not text:
            try:
                raw = page.get_text("rawdict") or {}
                spans: list[str] = []
                for block in raw.get("blocks", []):
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            t = span.get("text", "")
                            if t:
                                spans.append(t)
                text = " ".join(spans).strip()
            except Exception as exc:
                logger.debug("pymupdf rawdict page %s failed: %s", page.number, exc)

        # 4. html → strip tags
        if not text:
            try:
                html = page.get_text("html") or ""
                text = re.sub(r"<[^>]+>", " ", html).strip()
            except Exception as exc:
                logger.debug("pymupdf html page %s failed: %s", page.number, exc)

        page_texts.append(text)

    doc.close()
    return "\n".join(page_texts)


def _extract_pdf_pdfminer(data: bytes) -> str:
    """Extract text via pdfminer.six — handles CID/Type3 fonts that stump PyMuPDF."""
    from pdfminer.high_level import extract_text_to_fp
    from pdfminer.layout import LAParams

    out = io.StringIO()
    extract_text_to_fp(
        io.BytesIO(data),
        out,
        laparams=LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            boxes_flow=0.5,
            detect_vertical=False,
        ),
        output_type="text",
        codec="utf-8",
    )
    return out.getvalue()


def _clean(text: str) -> str:
    """Normalise whitespace without destroying paragraph structure."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove stray null bytes or other control chars that confuse the LLM
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text.strip()


def extract_pdf_text(data: bytes) -> str:
    """Full 5-strategy PDF extraction with pdfminer fallback."""
    # Strategies 1-4: PyMuPDF
    text = ""
    try:
        text = _extract_pdf_pymupdf(data)
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed entirely: %s", exc)

    # Strategy 5: pdfminer — only if pymupdf gave us nothing useful
    if len(text.strip()) < 50:
        logger.info("PyMuPDF returned <50 chars, falling back to pdfminer.six")
        try:
            text = _extract_pdf_pdfminer(data)
        except Exception as exc:
            logger.warning("pdfminer extraction also failed: %s", exc)

    return _clean(text)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def extract_docx_text(data: bytes) -> str:
    """Extract text from a DOCX including tables, headers, footers, and text boxes."""
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    # Body paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Tables (including nested)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    # Headers and footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for p in hf.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(t)

    # Text boxes / drawing canvas (wps:txbx or v:textbox XML elements)
    try:
        for node in doc.element.body.iter():
            tag = getattr(node, "tag", "") or ""
            if "txbx" in tag or "textbox" in tag.lower():
                box_texts = [
                    el.text
                    for el in node.iter()
                    if (getattr(el, "tag", "") or "").endswith("}t") and el.text
                ]
                if box_texts:
                    parts.append(" ".join(box_texts))
    except Exception as exc:
        logger.debug("DOCX text-box extraction failed: %s", exc)

    return _clean("\n".join(parts))


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(data: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf_text(data)
    if lower.endswith(".docx"):
        return extract_docx_text(data)
    raise UnsupportedFileError(f"Unsupported file type: {filename}")


def extract_and_scan(data: bytes, filename: str) -> tuple[str, object]:
    """Extract resume text and run integrity scanning.

    Returns (visible_text, IntegrityReport).
    Integrity scanning is best-effort — a crash there must never block extraction.
    """
    text = extract_text(data, filename)

    try:
        from .integrity import scan_resume_integrity
        report = scan_resume_integrity(data, filename)
    except Exception as exc:
        logger.warning("Integrity scan failed (non-fatal): %s", exc)
        # Return a safe empty report so the upload still succeeds
        from .integrity import IntegrityReport
        report = IntegrityReport(
            is_suspicious=False,
            severity="none",
            confidence=0.0,
            notes=["Integrity scan skipped due to an internal error."],
        )

    return text, report
